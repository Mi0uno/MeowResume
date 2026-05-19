"""Generate ATS-friendly graduate resumes as strict Word documents.

The template is intentionally conservative: A4, compact spacing, no text boxes,
reverse chronological entries, and quantified achievement bullets. The body
stays one-column and ATS-readable; an optional top-right ID photo slot is
available for China campus recruiting scenarios.
"""

from __future__ import annotations

import argparse
import copy
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

JSONDict = dict[str, Any]

BODY_EAST_ASIA_FONT = "DengXian"
BODY_LATIN_FONT = "Calibri"
HEADING_EAST_ASIA_FONT = "Microsoft YaHei"
HEADING_LATIN_FONT = "Calibri"
BODY_SIZE = 9.7
DETAIL_SIZE = 9
BULLET_SIZE = 9.2
NAME_SIZE = 17
SECTION_SIZE = 10.8
PHOTO_WIDTH_INCHES = 0.98
PHOTO_HEIGHT_INCHES = 1.38
HEADER_WIDTH_INCHES = 7.15
PHOTO_CELL_WIDTH_INCHES = 1.08
HEADER_TEXT_WIDTH_INCHES = HEADER_WIDTH_INCHES - (PHOTO_CELL_WIDTH_INCHES * 2)
ACTIVE_FORMAT: JSONDict = {}

SECTION_ORDER = (
    "education",
    "skills",
    "experience",
    "projects",
    "awards",
    "certifications",
)

SECTION_TITLES = {
    "summary": "求职摘要",
    "education": "教育经历",
    "skills": "专业技能",
    "experience": "实习经历",
    "projects": "项目&科研经历",
    "awards": "竞赛获奖",
    "certifications": "其它证书和荣誉",
}

DEFAULT_FORMAT: JSONDict = {
    "body_east_asia_font": BODY_EAST_ASIA_FONT,
    "body_latin_font": BODY_LATIN_FONT,
    "heading_east_asia_font": HEADING_EAST_ASIA_FONT,
    "heading_latin_font": HEADING_LATIN_FONT,
    "body_size": BODY_SIZE,
    "detail_size": DETAIL_SIZE,
    "bullet_size": BULLET_SIZE,
    "name_size": NAME_SIZE,
    "section_size": SECTION_SIZE,
    "line_spacing": 1.0,
    "top_margin": 0.45,
    "bottom_margin": 0.45,
    "left_margin": 0.55,
    "right_margin": 0.55,
}

TECH_KEYWORDS = [
    "Python",
    "Java",
    "Go",
    "C++",
    "C#",
    "JavaScript",
    "TypeScript",
    "SQL",
    "FastAPI",
    "Django",
    "Flask",
    "Spring Boot",
    "Spring",
    "React",
    "Vue",
    "Next.js",
    "Node.js",
    "Redis",
    "MySQL",
    "PostgreSQL",
    "MongoDB",
    "ClickHouse",
    "Kafka",
    "RabbitMQ",
    "Elasticsearch",
    "Docker",
    "Kubernetes",
    "Linux",
    "Git",
    "CI/CD",
    "AWS",
    "Azure",
    "GCP",
    "微服务",
    "分布式",
    "高并发",
    "网络安全",
    "渗透测试",
    "CTF",
    "机器学习",
    "深度学习",
    "大模型",
]

METRIC_PATTERN = re.compile(
    r"(?:P\d{2}|QPS|TPS)|"
    r"(?:\d+(?:\.\d+)?\+?\s*(?:%|％|个|项|台|次|万|亿|人|名|ms|s|秒|分钟|小时|天|倍|条|行|"
    r"MB|GB|TB|K|k|M|元|分|页|套|类|种))",
    re.IGNORECASE,
)

WINDOWS_UNSAFE_CHARS = re.compile(r'[<>:"/\\|?*\r\n\t]+')


def load_resume_data(path: Path) -> JSONDict:
    """Load a UTF-8 JSON resume file."""
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise FileNotFoundError(f"Resume data file not found: {path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in {path}: {exc}") from exc


def contains_metric(text: str) -> bool:
    """Return whether a bullet contains a measurable result or scale."""
    return bool(METRIC_PATTERN.search(text))


def extract_keywords(jd_text: str) -> list[str]:
    """Extract known technical keywords from a job description in reading order."""
    matches: list[tuple[int, str]] = []
    for keyword in TECH_KEYWORDS:
        index = find_keyword_index(jd_text, keyword)
        if index >= 0:
            matches.append((index, keyword))

    seen: set[str] = set()
    ordered: list[str] = []
    for _, keyword in sorted(matches, key=lambda item: item[0]):
        key = keyword.lower()
        if key not in seen:
            ordered.append(keyword)
            seen.add(key)
    return ordered


def find_keyword_index(text: str, keyword: str) -> int:
    """Find a keyword without matching it inside a longer ASCII identifier."""
    if re.search(r"[A-Za-z0-9]", keyword):
        pattern = re.compile(
            rf"(?<![A-Za-z0-9_.+#]){re.escape(keyword)}(?![A-Za-z0-9_.+#])",
            re.IGNORECASE,
        )
        match = pattern.search(text)
        return match.start() if match else -1
    return text.find(keyword)


def analyze_resume(data: JSONDict) -> list[str]:
    """Return practical warnings before generating the final resume."""
    warnings: list[str] = []
    basics = data.get("basics", {})
    for field in ("name", "phone", "email", "target_role"):
        if not basics.get(field):
            warnings.append(f"基础信息缺少 {field}，大厂 ATS 解析和 HR 联系都会受影响。")

    if not data.get("education"):
        warnings.append("应届生简历必须保留教育经历，并建议放在最前面。")
    if not data.get("skills"):
        warnings.append("缺少专业技能模块，技术岗建议用关键词对齐岗位 JD。")

    bullets = list(iter_bullets(data))
    if bullets:
        metric_count = sum(1 for bullet in bullets if contains_metric(bullet))
        if metric_count < max(1, len(bullets) // 2):
            warnings.append("项目/实习 bullet 的量化结果偏少，建议至少一半 bullet 带数字、规模、性能或业务结果。")
        for bullet in bullets:
            length = len(bullet)
            if not contains_metric(bullet):
                warnings.append(f"bullet 缺少量化结果，建议补充数字、规模、性能或结果：{bullet}")
            if length < 16 and not contains_metric(bullet):
                warnings.append(f"bullet 过短，像职责而不是成果：{bullet}")
            if length > 95:
                warnings.append(f"bullet 过长，Word 中容易换行拥挤：{bullet[:42]}...")
    else:
        warnings.append("缺少项目或实习 bullet，应届技术岗至少准备 6-10 条成果描述。")

    total_text = len(json.dumps(data, ensure_ascii=False))
    if len(bullets) > 12 or total_text > 3600:
        warnings.append("内容可能超过一页，建议优先保留目标岗位最相关的 2-3 段经历。")

    return warnings


def generate_resume_docx(data: JSONDict, output_path: Path, jd_text: str | None = None) -> None:
    """Generate a resume Word document at output_path."""
    global ACTIVE_FORMAT
    resume_data = copy.deepcopy(data)
    keywords = extract_keywords(jd_text or "")
    if keywords:
        resume_data["skills"] = reorder_skills_for_keywords(resume_data.get("skills"), keywords)

    previous_format = ACTIVE_FORMAT
    ACTIVE_FORMAT = resolve_format(resume_data.get("format"))
    document = Document()
    try:
        configure_document(document)

        add_header(document, resume_data.get("basics", {}))
        if resume_data.get("basics", {}).get("summary"):
            add_summary(document, resume_data["basics"]["summary"])

        for section in SECTION_ORDER:
            value = resume_data.get(section)
            if not value:
                continue
            if section == "education":
                add_education(document, value)
            elif section == "skills":
                add_skills(document, value, keywords)
            elif section == "experience":
                add_experience(document, value)
            elif section == "projects":
                add_projects(document, value)
            elif section == "awards":
                add_simple_list_section(document, SECTION_TITLES[section], value)
            elif section == "certifications":
                add_simple_list_section(document, SECTION_TITLES[section], value)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
    finally:
        ACTIVE_FORMAT = previous_format


def resolve_format(raw_format: Any) -> JSONDict:
    """Merge user-facing format controls with conservative defaults."""
    settings = copy.deepcopy(DEFAULT_FORMAT)
    if not isinstance(raw_format, dict):
        return settings

    text_fields = (
        "body_east_asia_font",
        "body_latin_font",
        "heading_east_asia_font",
        "heading_latin_font",
    )
    for field in text_fields:
        value = str(raw_format.get(field, "")).strip()
        if value:
            settings[field] = value[:40]

    ranges = {
        "body_size": (8.0, 12.0),
        "detail_size": (7.5, 11.0),
        "bullet_size": (8.0, 11.5),
        "name_size": (14.0, 24.0),
        "section_size": (9.5, 14.0),
        "line_spacing": (0.9, 1.35),
        "top_margin": (0.25, 0.9),
        "bottom_margin": (0.25, 0.9),
        "left_margin": (0.35, 0.9),
        "right_margin": (0.35, 0.9),
    }
    for field, bounds in ranges.items():
        settings[field] = clamp_float(raw_format.get(field), settings[field], *bounds)
    return settings


def clamp_float(value: Any, default: float, minimum: float, maximum: float) -> float:
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        return default
    return max(minimum, min(parsed, maximum))


def fmt(key: str) -> Any:
    return ACTIVE_FORMAT.get(key, DEFAULT_FORMAT[key])


def build_output_filename(data: JSONDict) -> str:
    """Build a Windows-safe resume filename."""
    basics = data.get("basics", {})
    school = ""
    education = data.get("education") or []
    if education:
        school = education[0].get("school", "")
    parts = [
        basics.get("name", "resume"),
        basics.get("target_role", "岗位"),
        school,
        basics.get("phone", ""),
    ]
    raw = "_".join(part for part in parts if part)
    safe = WINDOWS_UNSAFE_CHARS.sub("_", raw)
    safe = re.sub(r"\s+", "", safe).strip("._")
    return f"{safe or 'resume'}.docx"


def iter_bullets(data: JSONDict) -> list[str]:
    bullets: list[str] = []
    for section in ("experience", "projects"):
        for entry in data.get(section, []) or []:
            bullets.extend(str(item) for item in entry.get("bullets", []) or [])
            bullets.extend(str(item) for item in entry.get("results", []) or [])
    return bullets


def reorder_skills_for_keywords(skills: Any, keywords: list[str]) -> Any:
    """Move JD-matched skills to the front while preserving categories."""
    if not keywords or not skills:
        return skills

    keyword_map = {keyword.lower(): index for index, keyword in enumerate(keywords)}

    def priority(item: str) -> tuple[int, str]:
        lowered = item.lower()
        for key, index in keyword_map.items():
            if key in lowered:
                return (index, item)
        return (len(keyword_map) + 1, item)

    if isinstance(skills, dict):
        reordered: dict[str, list[str]] = {}
        for category, values in skills.items():
            values_list = [str(value) for value in coerce_list(values)]
            reordered[str(category)] = sorted(values_list, key=priority)
        return reordered

    return sorted([str(value) for value in coerce_list(skills)], key=priority)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(fmt("top_margin"))
    section.bottom_margin = Inches(fmt("bottom_margin"))
    section.left_margin = Inches(fmt("left_margin"))
    section.right_margin = Inches(fmt("right_margin"))

    normal = document.styles["Normal"]
    set_style_font(normal, size=fmt("body_size"))
    normal.paragraph_format.line_spacing = fmt("line_spacing")
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)


def add_header(document: Document, basics: JSONDict) -> None:
    if should_render_photo(basics):
        add_header_with_photo(document, basics)
        return

    name = basics.get("name", "姓名")
    role = basics.get("target_role", "")

    paragraph = document.add_paragraph()
    compact(paragraph, after=1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(name)
    set_run_font(run, size=fmt("name_size"), bold=True, east_asia=fmt("heading_east_asia_font"))

    contact_parts = [
        basics.get("phone", ""),
        basics.get("email", ""),
        basics.get("city", ""),
        role,
    ]
    add_centered_line(document, " | ".join(part for part in contact_parts if part), size=9)

    links = basics.get("links") or []
    link_parts = [
        f"{item.get('label', 'Link')}: {item.get('url', '')}"
        for item in links
        if item.get("url")
    ]
    if link_parts:
        add_centered_line(document, " | ".join(link_parts), size=8.5)


def add_header_with_photo(document: Document, basics: JSONDict) -> None:
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_width(table, HEADER_WIDTH_INCHES)
    set_fixed_table_layout(table)
    widths = [PHOTO_CELL_WIDTH_INCHES, HEADER_TEXT_WIDTH_INCHES, PHOTO_CELL_WIDTH_INCHES]
    set_table_grid(table, widths)
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)

    row = table.rows[0]
    row.height = Inches(PHOTO_HEIGHT_INCHES)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    spacer_cell, text_cell, photo_cell = row.cells
    set_cell_width(spacer_cell, PHOTO_CELL_WIDTH_INCHES)
    set_cell_width(text_cell, HEADER_TEXT_WIDTH_INCHES)
    set_cell_width(photo_cell, PHOTO_CELL_WIDTH_INCHES)
    set_cell_margins(spacer_cell, top=0, bottom=0, left=0, right=0)
    set_cell_margins(text_cell, top=0, bottom=0, left=0.03, right=0.03)
    set_cell_margins(photo_cell, top=0, bottom=0, left=0.04, right=0)
    spacer_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    clear_cell(spacer_cell)
    clear_cell(text_cell)
    clear_cell(photo_cell)
    add_header_text_to_cell(text_cell, basics)
    add_photo_to_cell(photo_cell, get_photo_path(basics))


def add_header_text_to_cell(cell: Any, basics: JSONDict) -> None:
    name = basics.get("name", "姓名")
    role = basics.get("target_role", "")

    paragraph = cell.paragraphs[0]
    compact(paragraph, after=1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        paragraph,
        name,
        size=fmt("name_size"),
        bold=True,
        east_asia=fmt("heading_east_asia_font"),
        latin=fmt("heading_latin_font"),
    )

    contact_parts = [
        basics.get("phone", ""),
        basics.get("email", ""),
        basics.get("city", ""),
        role,
    ]
    add_cell_centered_line(cell, " | ".join(part for part in contact_parts if part), size=9)

    links = basics.get("links") or []
    link_parts = [
        f"{item.get('label', 'Link')}: {item.get('url', '')}"
        for item in links
        if item.get("url")
    ]
    if link_parts:
        add_cell_centered_line(cell, " | ".join(link_parts), size=8.5, color="555555")


def add_photo_to_cell(cell: Any, photo_path: Path | None) -> None:
    set_cell_border(cell, color="7A7A7A", size="6")

    paragraph = cell.paragraphs[0]
    compact(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if photo_path and photo_path.exists():
        run = paragraph.add_run()
        run.add_picture(
            str(photo_path),
            width=Inches(PHOTO_WIDTH_INCHES),
            height=Inches(PHOTO_HEIGHT_INCHES),
        )
        return

    add_run(paragraph, "证件照", size=8.5, color="666666")
    hint = cell.add_paragraph()
    compact(hint)
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(hint, "25mm x 35mm", size=7.2, color="888888")


def add_cell_centered_line(cell: Any, text: str, size: float, color: str = "000000") -> None:
    paragraph = cell.add_paragraph()
    compact(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, text, size=size, color=color)


def add_summary(document: Document, summary: str) -> None:
    add_section_heading(document, SECTION_TITLES["summary"])
    paragraph = document.add_paragraph()
    compact(paragraph)
    add_run(paragraph, str(summary), size=fmt("body_size"))


def add_education(document: Document, entries: list[JSONDict]) -> None:
    add_section_heading(document, SECTION_TITLES["education"])
    for entry in entries:
        title = join_nonempty(
            [
                entry.get("school"),
                entry.get("major"),
                entry.get("degree"),
            ],
            " | ",
        )
        add_entry_line(document, title, format_period(entry), bold=True)

        details = []
        if entry.get("gpa"):
            details.append(f"GPA: {entry['gpa']}")
        if entry.get("rank"):
            details.append(f"排名: {entry['rank']}")
        if entry.get("courses"):
            details.append("相关课程: " + "、".join(coerce_list(entry["courses"])))
        if details:
            add_plain_line(document, " | ".join(details), size=fmt("detail_size"))

        honors = coerce_list(entry.get("honors"))
        if honors:
            add_plain_line(document, "荣誉: " + "、".join(honors), size=fmt("detail_size"))


def add_skills(document: Document, skills: Any, keywords: list[str] | None = None) -> None:
    add_section_heading(document, SECTION_TITLES["skills"])
    if isinstance(skills, dict):
        for category, values in skills.items():
            skill_text = " | ".join(str(value) for value in coerce_list(values))
            paragraph = document.add_paragraph()
            compact(paragraph)
            add_run(paragraph, f"{category}: ", size=fmt("detail_size"), bold=True)
            add_run(paragraph, skill_text, size=fmt("detail_size"))
    else:
        add_plain_line(document, " | ".join(str(value) for value in coerce_list(skills)), size=fmt("detail_size"))

    if keywords:
        add_plain_line(document, "JD 关键词: " + " | ".join(keywords[:10]), size=8.5, color="666666")


def add_experience(document: Document, entries: list[JSONDict]) -> None:
    add_section_heading(document, SECTION_TITLES["experience"])
    for entry in entries:
        title = join_nonempty(
            [
                entry.get("organization"),
                entry.get("role"),
                entry.get("location"),
            ],
            " | ",
        )
        add_entry_line(document, title, format_period(entry), bold=True)
        add_structured_entry_details(
            document,
            entry,
            content_label="工作内容",
            responsibilities_label="负责工作",
            results_label="工作成果",
        )


def add_projects(document: Document, entries: list[JSONDict]) -> None:
    add_section_heading(document, SECTION_TITLES["projects"])
    for entry in entries:
        tech = " / ".join(coerce_list(entry.get("technologies")))
        title = join_nonempty([entry.get("name"), entry.get("role"), tech], " | ")
        add_entry_line(document, title, format_period(entry), bold=True)
        add_structured_entry_details(
            document,
            entry,
            content_label="内容",
            responsibilities_label="负责工作",
            results_label="所获成果",
        )


def add_simple_list_section(document: Document, title: str, items: list[Any]) -> None:
    add_section_heading(document, title)
    for item in coerce_list(items):
        if isinstance(item, dict):
            left = join_nonempty([item.get("name"), item.get("issuer")], " | ")
            right = join_nonempty([item.get("level"), item.get("date")], " | ")
            add_compact_item_line(document, left, right, bold=True)
            add_structured_entry_details(
                document,
                item,
                content_label="内容",
                responsibilities_label="负责工作",
                results_label="所获成果",
            )
        else:
            left, right = split_compact_item(str(item))
            add_compact_item_line(document, left, right)


def split_compact_item(text: str) -> tuple[str, str]:
    parts = [part.strip() for part in re.split(r"\s*[|｜]\s*", text) if part.strip()]
    if len(parts) <= 1:
        return text.strip(), ""
    return parts[0], " | ".join(parts[1:])


def add_compact_item_line(document: Document, left: str, right: str = "", bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    compact(paragraph, before=0.2)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_TAB_ALIGNMENT.RIGHT)
    add_run(paragraph, left, size=fmt("detail_size"), bold=bold)
    if right:
        add_run(paragraph, "\t" + right, size=fmt("detail_size"), color="444444")


def add_structured_entry_details(
    document: Document,
    entry: JSONDict,
    content_label: str,
    responsibilities_label: str,
    results_label: str,
) -> None:
    content = str(entry.get("content", "")).strip()
    if content:
        add_labeled_line(document, content_label, content)

    responsibilities = coerce_list(entry.get("responsibilities"))
    if responsibilities:
        add_labeled_items(document, responsibilities_label, responsibilities)

    results = coerce_list(entry.get("results"))
    if results:
        add_labeled_items(document, results_label, results)

    if not content and not responsibilities and not results:
        add_bullets(document, entry.get("bullets", []))
    elif entry.get("bullets"):
        add_labeled_items(document, "补充", coerce_list(entry.get("bullets")))


def add_labeled_line(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    compact(paragraph, left=0.18, first_line=-0.18)
    add_run(paragraph, "• ", size=fmt("bullet_size"))
    add_run(paragraph, f"{label}：", size=fmt("bullet_size"), bold=True)
    add_run(paragraph, text, size=fmt("bullet_size"))


def add_labeled_items(document: Document, label: str, items: list[Any]) -> None:
    values = [str(item).strip() for item in items if str(item).strip()]
    if not values:
        return
    if len(values) == 1:
        add_labeled_line(document, label, values[0])
        return

    paragraph = document.add_paragraph()
    compact(paragraph, left=0.18, first_line=-0.18)
    add_run(paragraph, "• ", size=fmt("bullet_size"))
    add_run(paragraph, f"{label}：", size=fmt("bullet_size"), bold=True)
    for item in values:
        sub = document.add_paragraph()
        compact(sub, left=0.42, first_line=-0.16)
        add_run(sub, "• ", size=fmt("bullet_size"))
        add_run(sub, item, size=fmt("bullet_size"))


def add_bullets(document: Document, bullets: list[str]) -> None:
    for bullet in coerce_list(bullets):
        paragraph = document.add_paragraph()
        compact(paragraph, left=0.18, first_line=-0.18)
        add_run(paragraph, "• ", size=fmt("bullet_size"))
        add_run(paragraph, str(bullet), size=fmt("bullet_size"))


def add_entry_line(document: Document, left: str, right: str = "", bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    compact(paragraph, before=0.4)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_TAB_ALIGNMENT.RIGHT)
    add_run(paragraph, left, size=fmt("body_size"), bold=bold)
    if right:
        add_run(paragraph, "\t" + right, size=fmt("detail_size"), bold=bold)


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    compact(paragraph, before=3.5, after=1)
    add_bottom_border(paragraph)
    run = paragraph.add_run(title)
    set_run_font(run, size=fmt("section_size"), bold=True, east_asia=fmt("heading_east_asia_font"))


def add_centered_line(document: Document, text: str, size: float = 8.5) -> None:
    paragraph = document.add_paragraph()
    compact(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, text, size=size)


def add_plain_line(
    document: Document,
    text: str,
    size: float = 9,
    color: str = "000000",
    bold: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    compact(paragraph)
    add_run(paragraph, text, size=size, color=color, bold=bold)


def add_run(
    paragraph: Any,
    text: str,
    size: float,
    bold: bool = False,
    color: str = "000000",
    east_asia: str | None = None,
    latin: str | None = None,
) -> Any:
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        bold=bold,
        color=color,
        east_asia=east_asia or fmt("body_east_asia_font"),
        latin=latin or fmt("body_latin_font"),
    )
    return run


def compact(
    paragraph: Any,
    before: float = 0,
    after: float = 0,
    left: float = 0,
    first_line: float = 0,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = fmt("line_spacing")
    if left:
        paragraph.paragraph_format.left_indent = Inches(left)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(first_line)


def add_bottom_border(paragraph: Any, color: str = "7A7A7A") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_width(table: Any, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_grid(table: Any, widths_inches: list[float]) -> None:
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)

    for grid_col in list(tbl_grid.findall(qn("w:gridCol"))):
        tbl_grid.remove(grid_col)

    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)


def clear_cell(cell: Any) -> None:
    cell.text = ""


def set_cell_width(cell: Any, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def get_cell_width_inches(cell: Any) -> float:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        return 0.0
    raw_width = tc_w.get(qn("w:w"))
    if not raw_width:
        return 0.0
    return int(raw_width) / 1440


def set_cell_margins(
    cell: Any,
    top: float = 0,
    bottom: float = 0,
    left: float = 0,
    right: float = 0,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    margins = {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }
    for side, width in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(int(width * 1440)))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell: Any, color: str = "7A7A7A", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_style_font(style: Any, size: float) -> None:
    style.font.name = fmt("body_latin_font")
    style.font.size = Pt(size)
    r_pr = style._element.get_or_add_rPr()
    ensure_rfonts(r_pr).set(qn("w:eastAsia"), fmt("body_east_asia_font"))
    ensure_rfonts(r_pr).set(qn("w:ascii"), fmt("body_latin_font"))
    ensure_rfonts(r_pr).set(qn("w:hAnsi"), fmt("body_latin_font"))


def set_run_font(
    run: Any,
    size: float,
    bold: bool = False,
    color: str = "000000",
    east_asia: str | None = None,
    latin: str | None = None,
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    effective_latin = latin or fmt("body_latin_font")
    effective_east_asia = east_asia or fmt("body_east_asia_font")
    run.font.name = effective_latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = ensure_rfonts(r_pr)
    r_fonts.set(qn("w:eastAsia"), effective_east_asia)
    r_fonts.set(qn("w:ascii"), effective_latin)
    r_fonts.set(qn("w:hAnsi"), effective_latin)


def ensure_rfonts(r_pr: Any) -> Any:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    return r_fonts


def format_period(entry: JSONDict) -> str:
    start = str(entry.get("start", "")).strip()
    end = str(entry.get("end", "")).strip()
    if start and end:
        return f"{start} - {end}"
    return start or end


def join_nonempty(parts: list[Any], sep: str) -> str:
    return sep.join(str(part).strip() for part in parts if str(part or "").strip())


def coerce_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def should_render_photo(basics: JSONDict) -> bool:
    photo = basics.get("photo")
    if isinstance(photo, dict):
        return bool(photo.get("enabled", bool(photo.get("path"))))
    if isinstance(photo, str):
        return bool(photo.strip())
    return False


def get_photo_path(basics: JSONDict) -> Path | None:
    photo = basics.get("photo")
    if isinstance(photo, dict):
        raw_path = str(photo.get("path", "")).strip()
    elif isinstance(photo, str):
        raw_path = photo.strip()
    else:
        raw_path = ""
    if not raw_path:
        return None
    return Path(raw_path).expanduser()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a strict Word resume from JSON.")
    parser.add_argument("--data", type=Path, default=Path("candidate.example.json"), help="Resume JSON file.")
    parser.add_argument("--output", type=Path, help="Output .docx path. Defaults to output/<name_role_school_phone>.docx.")
    parser.add_argument("--jd", type=Path, help="Optional job description text file for keyword prioritization.")
    parser.add_argument("--photo", type=Path, help="Optional ID photo path. Overrides basics.photo.path.")
    parser.add_argument("--no-photo", action="store_true", help="Disable the top-right ID photo slot.")
    parser.add_argument("--strict", action="store_true", help="Exit with code 2 if resume quality warnings are found.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    data = load_resume_data(args.data)
    if args.photo:
        data.setdefault("basics", {})["photo"] = {"enabled": True, "path": str(args.photo)}
    if args.no_photo:
        data.setdefault("basics", {})["photo"] = {"enabled": False, "path": ""}
    jd_text = args.jd.read_text(encoding="utf-8") if args.jd else None
    output = args.output or Path("output") / build_output_filename(data)

    warnings = analyze_resume(data)
    generate_resume_docx(data, output, jd_text=jd_text)

    print(f"已生成 Word 简历: {output.resolve()}")
    if warnings:
        print("\n简历质量提醒:")
        for warning in warnings:
            print(f"- {warning}")
    if args.strict and warnings:
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

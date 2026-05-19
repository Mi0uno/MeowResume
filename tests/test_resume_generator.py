import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

import resume_generator as rg


def collect_document_text(document):
    text = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text.extend(p.text for p in cell.paragraphs)
    return "\n".join(text)


class ResumeGeneratorTests(unittest.TestCase):
    def test_contains_metric_detects_numbers_percentages_and_chinese_units(self):
        self.assertTrue(rg.contains_metric("将接口 P95 延迟降低 38%，覆盖 12 个核心场景"))
        self.assertTrue(rg.contains_metric("支撑日均 10 万次请求，错误率低于 0.1%"))
        self.assertTrue(rg.contains_metric("维护 3 台服务器与 2 个业务系统"))
        self.assertFalse(rg.contains_metric("负责后端接口开发和数据库设计"))

    def test_extract_keywords_prefers_technical_terms(self):
        jd = "岗位要求：熟悉 Python、FastAPI、Redis、MySQL，了解 Docker 和 Linux。"
        keywords = rg.extract_keywords(jd)
        self.assertEqual(keywords[:5], ["Python", "FastAPI", "Redis", "MySQL", "Docker"])

    def test_analyze_resume_warns_when_bullets_are_not_quantified(self):
        data = rg.load_resume_data(Path("candidate.example.json"))
        data["projects"][0]["bullets"] = ["负责系统开发，参与需求分析"]

        warnings = rg.analyze_resume(data)

        self.assertTrue(any("量化" in warning for warning in warnings))

    def test_generate_docx_contains_core_sections(self):
        data = rg.load_resume_data(Path("candidate.example.json"))

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "resume.docx"
            rg.generate_resume_docx(data, output)
            self.assertTrue(output.exists())

            document = Document(output)
            text = collect_document_text(document)
            self.assertIn(data["basics"]["name"], text)
            self.assertIn("教育经历", text)
            self.assertIn("项目&科研经历", text)
            self.assertIn("竞赛获奖", text)
            self.assertIn("其它证书和荣誉", text)
            self.assertIn("专业技能", text)

    def test_generate_docx_renders_structured_experience_labels(self):
        data = rg.load_resume_data(Path("candidate.example.json"))
        data["projects"][0]["content"] = "针对机器人自主移动任务进行路径规划与感知优化"
        data["projects"][0]["responsibilities"] = [
            "负责导航模块设计，完成路径规划与避障策略调优",
            "负责实验记录整理与论文初稿撰写"
        ]
        data["projects"][0]["results"] = [
            "识别准确率提升至 92%，路径规划耗时降低 28%"
        ]
        data["awards"] = [
            {
                "name": "RoboCup 机器人世界杯中国赛",
                "level": "国家级二等奖",
                "date": "2025.08",
                "responsibilities": [
                    "负责机器人软件部分，包括 STM32 电机控制与多传感器数据融合"
                ],
                "results": [
                    "颜色识别准确率和比赛稳定性显著提升"
                ]
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "resume_structured.docx"
            rg.generate_resume_docx(data, output)
            text = collect_document_text(Document(output))

        self.assertIn("内容：针对机器人自主移动任务进行路径规划与感知优化", text)
        self.assertIn("负责工作：", text)
        self.assertIn("所获成果：", text)
        self.assertIn("RoboCup 机器人世界杯中国赛", text)
        self.assertIn("国家级二等奖", text)

    def test_split_compact_item_moves_award_metadata_right(self):
        left, right = rg.split_compact_item("2025 全国大学生信息安全竞赛区域赛 | 二等奖 | 2025")

        self.assertEqual(left, "2025 全国大学生信息安全竞赛区域赛")
        self.assertEqual(right, "二等奖 | 2025")

    def test_generate_docx_contains_photo_placeholder_when_enabled(self):
        data = rg.load_resume_data(Path("candidate.example.json"))
        data["basics"]["photo"] = {"enabled": True, "path": ""}

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "resume_with_photo.docx"
            rg.generate_resume_docx(data, output)

            document = Document(output)
            self.assertGreaterEqual(len(document.tables), 1)
            self.assertIn("证件照", collect_document_text(document))
            cells = document.tables[0].rows[0].cells
            self.assertEqual(len(cells), 3)
            self.assertLessEqual(rg.get_cell_width_inches(cells[2]), 1.25)
            self.assertAlmostEqual(
                rg.get_cell_width_inches(cells[0]),
                rg.get_cell_width_inches(cells[2]),
                places=2,
            )
            self.assertIn(data["basics"]["name"], "\n".join(p.text for p in cells[1].paragraphs))

    def test_generate_docx_can_disable_photo_placeholder(self):
        data = rg.load_resume_data(Path("candidate.example.json"))
        data["basics"]["photo"] = {"enabled": False, "path": ""}

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "resume_without_photo.docx"
            rg.generate_resume_docx(data, output)

            document = Document(output)
            self.assertNotIn("证件照", collect_document_text(document))

    def test_generate_docx_applies_format_settings(self):
        data = rg.load_resume_data(Path("candidate.example.json"))
        data["basics"]["photo"] = {"enabled": False, "path": ""}
        data["format"] = {
            "body_east_asia_font": "SimSun",
            "body_latin_font": "Times New Roman",
            "heading_east_asia_font": "SimHei",
            "heading_latin_font": "Arial",
            "body_size": 10.5,
            "name_size": 20,
            "section_size": 12,
            "line_spacing": 1.2,
            "top_margin": 0.6,
            "bottom_margin": 0.65,
            "left_margin": 0.7,
            "right_margin": 0.75,
        }

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "resume_format.docx"
            rg.generate_resume_docx(data, output)
            document = Document(output)

        section = document.sections[0]
        self.assertAlmostEqual(section.top_margin.inches, 0.6, places=2)
        self.assertAlmostEqual(section.bottom_margin.inches, 0.65, places=2)
        self.assertAlmostEqual(section.left_margin.inches, 0.7, places=2)
        self.assertAlmostEqual(section.right_margin.inches, 0.75, places=2)
        self.assertAlmostEqual(document.styles["Normal"].font.size.pt, 10.5, places=1)
        self.assertAlmostEqual(document.styles["Normal"].paragraph_format.line_spacing, 1.2, places=1)

        text_runs = [
            run
            for paragraph in document.paragraphs
            for run in paragraph.runs
            if run.text == data["basics"]["name"]
        ]
        self.assertTrue(text_runs)
        self.assertAlmostEqual(text_runs[0].font.size.pt, 20, places=1)

    def test_build_output_filename_is_windows_safe(self):
        data = json.loads(Path("candidate.example.json").read_text(encoding="utf-8"))
        filename = rg.build_output_filename(data)

        self.assertTrue(filename.endswith(".docx"))
        self.assertNotIn("/", filename)
        self.assertNotIn("\\", filename)


if __name__ == "__main__":
    unittest.main()
